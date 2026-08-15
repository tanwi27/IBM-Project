import hashlib
import os
import pdfplumber
from docx import Document
from typing import Dict, Tuple, Any

def calculate_sha256(text: str) -> str:
    """Computes a SHA-256 hash of the text to enable deterministic score caching."""
    return hashlib.sha256(text.encode('utf-8')).hexdigest()

def detect_columns_pdf(page) -> bool:
    """
    Detects if a PDF page uses a multi-column layout.
    Analyzes horizontal spacing of characters on the page.
    """
    words = page.extract_words()
    if not words:
        return False
    
    # Divide the page width into intervals and count word density
    page_width = float(page.width)
    left_margin = page_width * 0.15
    right_margin = page_width * 0.85
    center_line = page_width / 2.0
    gap_width = page_width * 0.08 # column gap of 8% of page width
    
    col1_count = 0
    col2_count = 0
    center_gap_count = 0
    
    for w in words:
        x0 = float(w['x0'])
        x1 = float(w['x1'])
        
        # Word is in the left column band
        if x1 < (center_line - gap_width / 2) and x0 > left_margin:
            col1_count += 1
        # Word is in the right column band
        elif x0 > (center_line + gap_width / 2) and x1 < right_margin:
            col2_count += 1
        # Word is crossing the center gap
        elif x0 < center_line and x1 > center_line:
            center_gap_count += 1

    # If there are healthy counts of words in left and right bands, and low count in center gap,
    # it indicates a multi-column layout.
    total_words = len(words)
    if total_words > 50:
        col_ratio = (col1_count + col2_count) / total_words
        gap_ratio = center_gap_count / total_words
        if col_ratio > 0.6 and gap_ratio < 0.1:
            return True
            
    return False

def parse_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Parses a PDF file. Extracts text and detects structural elements.
    """
    text_content = []
    structural_metadata = {
        "has_tables": False,
        "tables_count": 0,
        "has_images": False,
        "images_count": 0,
        "has_columns": False,
        "has_text_boxes": False, # Text boxes in PDFs are often parsed as drawings or separate blocks
        "page_count": 0
    }
    
    with pdfplumber.open(file_path) as pdf:
        structural_metadata["page_count"] = len(pdf.pages)
        
        for i, page in enumerate(pdf.pages):
            # Extract plain text
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
                
            # Tables
            tables = page.find_tables()
            if tables:
                structural_metadata["has_tables"] = True
                structural_metadata["tables_count"] += len(tables)
                
            # Images
            images = page.images
            if images:
                structural_metadata["has_images"] = True
                structural_metadata["images_count"] += len(images)
                
            # Column layout
            if not structural_metadata["has_columns"]:
                if detect_columns_pdf(page):
                    structural_metadata["has_columns"] = True

    combined_text = "\n".join(text_content)
    return combined_text, structural_metadata

def parse_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Parses a DOCX file. Extracts text and detects structural elements.
    """
    doc = Document(file_path)
    text_content = []
    
    # Read paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_content.append(paragraph.text)
            
    structural_metadata = {
        "has_tables": False,
        "tables_count": 0,
        "has_images": False,
        "images_count": 0,
        "has_columns": False,
        "has_text_boxes": False,
        "page_count": 0 # Exact page count for DOCX requires page break analysis or is estimated
    }
    
    # Detect Tables
    if len(doc.tables) > 0:
        structural_metadata["has_tables"] = True
        structural_metadata["tables_count"] = len(doc.tables)
        # Extract table text as well so it's not lost to parsing
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text_content.append(" | ".join(row_text))

    # Detect Columns
    for section in doc.sections:
        # Check if the section has more than 1 column
        if section.columns and len(section.columns) > 1:
            structural_metadata["has_columns"] = True
            break
            
    # Detect Images / Shapes
    # Inline shapes are images, text boxes, or charts
    inline_shapes = doc.inline_shapes
    if len(inline_shapes) > 0:
        structural_metadata["has_images"] = True
        structural_metadata["images_count"] = len(inline_shapes)
        
    # Check for text boxes. Text boxes are floating shapes, which reside in document part relations
    # Let's inspect document body XML for text boxes (w:txbxContent)
    try:
        xml_content = doc.element.xml
        if "w:txbxContent" in xml_content:
            structural_metadata["has_text_boxes"] = True
    except Exception:
        pass

    # Word count estimation for pages (approx 450 words per page in a resume format)
    combined_text = "\n".join(text_content)
    word_count = len(combined_text.split())
    structural_metadata["page_count"] = max(1, round(word_count / 450))
    
    return combined_text, structural_metadata

def parse_resume_file(file_path: str) -> Tuple[str, str, Dict[str, Any]]:
    """
    Main entrypoint for parsing a resume. Handles file type validation, text extraction,
    and structural metadata compilation.
    """
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        text, metadata = parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        text, metadata = parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}. Only PDF and DOCX are supported.")
    
    file_hash = calculate_sha256(text)
    return text, file_hash, metadata
